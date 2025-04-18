from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from database import Database
from datetime import datetime
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart

class ReportGenerator:
    def __init__(self):
        self.db = Database()
        self.setup_fonts()
        self.smtp_settings = None
        self.verify_smtp_settings()
    
    def verify_smtp_settings(self):
        return False
    
    def setup_fonts(self):
        # تسجيل الخطوط العربية
        try:
            # تحديد مسار مجلد الخطوط
            fonts_dir = os.path.join(os.path.dirname(__file__), 'fonts')
            
            # تسجيل الخطوط العربية مع التأكد من وجودها
            arabic_fonts = {
                'Cairo': 'Cairo-Regular.ttf',
                'CairoBold': 'Cairo-Bold.ttf',
                'NotoNaskhArabic': 'NotoNaskhArabic-Regular.ttf',
                'NotoNaskhArabicBold': 'NotoNaskhArabic-Bold.ttf',
                'Amiri': 'Amiri-Regular.ttf',
                'AmiriBold': 'Amiri-Bold.ttf'
            }
            
            for font_name, font_file in arabic_fonts.items():
                font_path = os.path.join(fonts_dir, font_file)
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                else:
                    print(f'تحذير: ملف الخط {font_file} غير موجود')
            
            # تسجيل عائلات الخطوط
            pdfmetrics.registerFontFamily('Cairo', normal='Cairo', bold='CairoBold')
            pdfmetrics.registerFontFamily('NotoNaskhArabic', normal='NotoNaskhArabic', bold='NotoNaskhArabicBold')
            pdfmetrics.registerFontFamily('Amiri', normal='Amiri', bold='AmiriBold')
            
            print('تم تحميل الخطوط العربية بنجاح')
        except Exception as e:
            print(f'خطأ في تحميل الخطوط العربية: {str(e)}')
            print('سيتم استخدام الخط الافتراضي')
            # تسجيل خط بديل في حالة الفشل
            pdfmetrics.registerFont(TTFont('Cairo', os.path.join(fonts_dir, 'Cairo-Regular.ttf')))
            pdfmetrics.registerFontFamily('Cairo', normal='Cairo')
    
    def get_arabic_styles(self):
        styles = getSampleStyleSheet()
        # تعريف نمط العنوان
        styles.add(ParagraphStyle(
            name='Title',
            fontName='Cairo',
            fontSize=24,
            leading=30,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=colors.HexColor('#2C3E50'),
            borderPadding=10,
            borderColor=colors.HexColor('#BDC3C7'),
            borderWidth=1,
            backColor=colors.HexColor('#ECF0F1'),
            wordWrap='RTL'
        ))
        
        # تعديل النمط العادي
        styles['Normal'] = ParagraphStyle(
            'Normal',
            fontName='NotoNaskhArabic',
            fontSize=14,
            leading=20,
            alignment=TA_RIGHT,
            textColor=colors.HexColor('#2C3E50'),
            wordWrap='RTL',
            firstLineIndent=20
        )
        
        # تعريف النمط الصغير
        styles.add(ParagraphStyle(
            name='Small',
            fontName='Amiri',
            fontSize=12,
            leading=16,
            alignment=TA_RIGHT,
            textColor=colors.HexColor('#7F8C8D'),
            wordWrap='RTL'
        ))
        
        # نمط خلايا الجدول
        styles.add(ParagraphStyle(
            name='TableCell',
            fontName='NotoNaskhArabic',
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#2C3E50'),
            wordWrap='RTL'
        ))
        return styles

    def create_pdf_report(self, report_type, date_filter=None, save_local=True):
        from arabic_styles import get_arabic_styles
        try:
            filename = f'reports/distribution_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            if not os.path.exists('reports'):
                os.makedirs('reports')
            
            # إعداد المستند مع هوامش مناسبة للغة العربية
            doc = SimpleDocTemplate(
                filename,
                pagesize=A4,
                rightMargin=40,
                leftMargin=40,
                topMargin=40,
                bottomMargin=40
            )
            
            elements = []
            styles = get_arabic_styles()
            
            # إضافة العنوان بتنسيق محسن
            title = Paragraph('تقرير توزيع المراقبين', styles['Title'])
            elements.append(title)
            
            if date_filter:
                date_info = Paragraph(f'تاريخ التقرير: {date_filter}', styles['Normal'])
                elements.append(date_info)
            
            elements.append(Spacer(1, 20))
            
            if report_type == 'distribution':
                data = self.get_distribution_data(date_filter)
                headers = ['التاريخ', 'القاعة', 'المراقب الأول', 'المراقب الثاني']
            elif report_type == 'statistics':
                data = self.get_statistics_data()
                headers = ['المراقب', 'عدد مرات المراقبة', 'القاعات']
            
            # تحويل البيانات إلى نص عربي مع تنسيق محسن
            table_data = [
                [Paragraph(str(cell), styles['TableCell' if i > 0 else 'Title']) for cell in row]
                for i, row in enumerate([headers] + list(data))
            ]
            
            # إنشاء الجدول مع تنسيقات محسنة للغة العربية
            col_widths = [doc.width/len(headers) for _ in headers]
            table = Table(table_data, colWidths=col_widths)
            
            table_style = [
                # تنسيق الرأس
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2C3E50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Cairo'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),
                
                # تنسيق البيانات
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#2C3E50')),
                ('FONTNAME', (0, 1), (-1, -1), 'NotoNaskhArabic'),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                
                # تحسين المسافات والهوامش
                ('TOPPADDING', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                
                # إضافة تأثيرات بصرية
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#34495E')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')])
            ]
            
            table.setStyle(TableStyle(table_style))
            elements.append(table)
            
            # إضافة التوقيع والتاريخ
            elements.append(Spacer(1, 30))
            signature = Paragraph(
                f'تم إنشاء التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
                styles['Small']
            )
            elements.append(signature)
            
            doc.build(elements)
            if save_local:
                print(f"تم حفظ التقرير بنجاح في: {filename}")
                return filename
            else:
                os.remove(filename)
                return None
        except Exception as e:
            print(f"حدث خطأ أثناء إنشاء التقرير: {str(e)}")
            if os.path.exists(filename):
                os.remove(filename)
            return None

    def send_report_by_email(self, filename, recipient_email):
        try:
            if not os.path.exists(filename):
                print("خطأ: لم يتم العثور على ملف التقرير")
                return False

            if not self.verify_smtp_settings():
                print("خطأ: يرجى التحقق من إعدادات SMTP وتأكد من إضافة كلمة مرور التطبيق من https://myaccount.google.com/apppasswords")
                return False

            msg = MIMEMultipart()
            msg['Subject'] = 'تقرير توزيع المراقبين'
            msg['From'] = self.smtp_settings['username']
            msg['To'] = recipient_email

            with open(filename, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='pdf' if filename.endswith('.pdf') else 'docx')
                attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(filename))
                msg.attach(attachment)

            with smtplib.SMTP(self.smtp_settings['host'], self.smtp_settings['port']) as server:
                server.starttls()
                server.login(self.smtp_settings['username'], self.smtp_settings['password'])
                server.send_message(msg)

            print("تم إرسال التقرير بنجاح عبر البريد الإلكتروني")
            return True

        except smtplib.SMTPAuthenticationError:
            print("خطأ في المصادقة: تأكد من صحة اسم المستخدم وكلمة المرور لحساب البريد الإلكتروني")
            print("يجب استخدام كلمة مرور التطبيق من إعدادات Google: https://myaccount.google.com/apppasswords")
            return False
        except smtplib.SMTPException as e:
            print(f"خطأ في إرسال البريد الإلكتروني: {str(e)}")
            return False
        except Exception as e:
            print(f"حدث خطأ غير متوقع: {str(e)}")
            return False
    
    def create_word_report(self, report_type, date_filter=None, custom_style=None):
        filename = f'reports/distribution_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        if not os.path.exists('reports'):
            os.makedirs('reports')
        
        doc = Document()
        
        # إعداد الخط الافتراضي للمستند
        doc.styles['Normal'].font.rtl = True
        doc.styles['Normal'].font.name = 'Traditional Arabic'
        doc.styles['Normal'].font.size = Pt(12)
        
        # إضافة العنوان الرئيسي مع تنسيق محسن
        heading = doc.add_heading('تقرير توزيع المراقبين', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.rtl = True
        heading.runs[0].font.name = 'Traditional Arabic'
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)  # #2C3E50
        
        # إضافة معلومات التقرير
        report_info = doc.add_paragraph()
        report_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        report_info.paragraph_format.space_after = Pt(20)
        
        if date_filter:
            date_run = report_info.add_run(f'تاريخ التقرير: {date_filter}')
            date_run.font.name = 'Traditional Arabic'
            date_run.font.size = Pt(12)
            date_run.font.rtl = True
        
        if report_type == 'distribution':
            data = self.get_distribution_data(date_filter)
            headers = ['التاريخ', 'القاعة', 'المراقب الأول', 'المراقب الثاني']
        elif report_type == 'statistics':
            data = self.get_statistics_data()
            headers = ['المراقب', 'عدد مرات المراقبة', 'القاعات']
        
        # إنشاء الجدول مع تنسيقات محسنة
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # تنسيق رؤوس الجدول
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.rtl = True
            
            # تطبيق تنسيق الخلفية للرؤوس
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # إضافة البيانات مع تنسيق محسن
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                cell = row_cells[i]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].runs[0].font.rtl = True
        
        # تطبيق تنسيقات إضافية للجدول
        for row in table.rows:
            row.height = Pt(30)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # إضافة التوقيع والتاريخ
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_run = signature.add_run(f'تم إنشاء التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        signature_run.font.name = 'Traditional Arabic'
        signature_run.font.size = Pt(10)
        signature_run.font.rtl = True
        
        try:
            doc.save(filename)
            print(f"تم حفظ التقرير بنجاح في: {filename}")
            return filename
        except Exception as e:
            print(f"حدث خطأ أثناء إنشاء التقرير: {str(e)}")
            if os.path.exists(filename):
                os.remove(filename)
            return None

    def send_report_by_email(self, filename, recipient_email):
        try:
            if not os.path.exists(filename):
                print("خطأ: لم يتم العثور على ملف التقرير")
                return False

            if not self.verify_smtp_settings():
                print("خطأ: فشل التحقق من إعدادات SMTP")
                return False

            msg = MIMEMultipart()
            msg['Subject'] = 'تقرير توزيع المراقبين'
            msg['From'] = self.smtp_settings['username']
            msg['To'] = recipient_email

            with open(filename, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='pdf' if filename.endswith('.pdf') else 'docx')
                attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(filename))
                msg.attach(attachment)

            with smtplib.SMTP(self.smtp_settings['host'], self.smtp_settings['port']) as server:
                server.starttls()
                server.login(self.smtp_settings['username'], self.smtp_settings['password'])
                server.send_message(msg)

            print("تم إرسال التقرير بنجاح عبر البريد الإلكتروني")
            return True

        except smtplib.SMTPAuthenticationError:
            print("خطأ في المصادقة: تأكد من صحة اسم المستخدم وكلمة المرور لحساب البريد الإلكتروني")
            print("تأكد من استخدام كلمة مرور التطبيق من إعدادات Google: https://myaccount.google.com/apppasswords")
            return False
        except smtplib.SMTPException as e:
            print(f"خطأ في إرسال البريد الإلكتروني: {str(e)}")
            return False
        except Exception as e:
            print(f"حدث خطأ غير متوقع: {str(e)}")
            return False
    
    def get_distribution_data(self, date_filter=None):
        cursor = self.db.conn.cursor()
        query = '''
            SELECT ed.date, r.name, t1.name, t2.name
            FROM distributions d
            JOIN exam_dates ed ON d.date_id = ed.id
            JOIN rooms r ON d.room_id = r.id
            JOIN teachers t1 ON d.teacher1_id = t1.id
            JOIN teachers t2 ON d.teacher2_id = t2.id
        '''
        
        if date_filter:
            query += ' WHERE ed.date = ?'
            cursor.execute(query, (date_filter,))
        else:
            cursor.execute(query)
        
        return cursor.fetchall()
    
    def get_statistics_data(self):
        cursor = self.db.conn.cursor()
        
        # إحصائيات المراقبين
        cursor.execute('''
            SELECT 
                t.name,
                COUNT(d.id) as supervision_count,
                GROUP_CONCAT(DISTINCT r.name) as rooms,
                ROUND(AVG(CASE WHEN d.teacher1_id = t.id OR d.teacher2_id = t.id THEN 1 ELSE 0 END) * 100, 2) as supervision_rate,
                COUNT(DISTINCT date(ed.date)) as unique_days,
                GROUP_CONCAT(DISTINCT strftime('%A', ed.date)) as weekdays
            FROM teachers t
            LEFT JOIN distributions d ON t.id = d.teacher1_id OR t.id = d.teacher2_id
            LEFT JOIN rooms r ON d.room_id = r.id
            LEFT JOIN exam_dates ed ON d.date_id = ed.id
            GROUP BY t.id
            ORDER BY supervision_count DESC
        ''')
        teacher_stats = cursor.fetchall()
        
        # إحصائيات القاعات
        cursor.execute('''
            SELECT 
                r.name,
                COUNT(d.id) as exam_count,
                r.capacity,
                GROUP_CONCAT(DISTINCT t.name) as supervisors
            FROM rooms r
            LEFT JOIN distributions d ON r.id = d.room_id
            LEFT JOIN teachers t ON d.teacher1_id = t.id OR d.teacher2_id = t.id
            GROUP BY r.id
            ORDER BY exam_count DESC
        ''')
        room_stats = cursor.fetchall()
        
        # إحصائيات الأيام
        cursor.execute('''
            SELECT 
                strftime('%A', ed.date) as weekday,
                COUNT(d.id) as exam_count,
                COUNT(DISTINCT r.id) as room_count,
                COUNT(DISTINCT t.id) as teacher_count
            FROM exam_dates ed
            LEFT JOIN distributions d ON ed.id = d.date_id
            LEFT JOIN rooms r ON d.room_id = r.id
            LEFT JOIN teachers t ON d.teacher1_id = t.id OR d.teacher2_id = t.id
            GROUP BY weekday
            ORDER BY exam_count DESC
        ''')
        day_stats = cursor.fetchall()
        
        return {
            'teachers': teacher_stats,
            'rooms': room_stats,
            'days': day_stats
        }

    def load_smtp_settings(self):
        """تحميل إعدادات SMTP من ملف تكوين"""
        try:
            with open('config.json', 'r') as f:
                config = json.load(f)
                self.smtp_settings = config.get('smtp', {})
                return True
        except Exception as e:
            print(f"خطأ في تحميل إعدادات SMTP: {str(e)}")
            return False

    def create_report(self, report_type, output_format='pdf', **kwargs):
        """إنشاء تقرير بتنسيق محدد"""
        try:
            if output_format.lower() == 'pdf':
                return self.create_pdf_report(report_type, **kwargs)
            elif output_format.lower() == 'docx':
                return self.create_word_report(report_type, **kwargs)
            else:
                raise ValueError("تنسيق التقرير غير مدعوم")
        except Exception as e:
            print(f"خطأ في إنشاء التقرير: {str(e)}")
            return None